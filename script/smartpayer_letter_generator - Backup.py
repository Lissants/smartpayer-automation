#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SmartPayer Letter Generator
============================

Summary of changes vs. the previous version
-------------------------------------------
1. Replaced the `letter_date_num` field with `year_num` (last two digits of
   the year). The template placeholder "Date Number (i.e., 26)" now maps
   to `year_num` so the rendered Letter No. ends with `…/V/26` instead of
   the day-of-month.
2. The `letter_number` field's last 3 digits (`nnn`) are now a dynamic
   per-batch sequence counter. The GUI / watcher passes a `--batch-token`
   on every run; the first file in a batch gets `…001`, the second `…002`,
   etc. The counter is stored in `.smartpayer_batch_state.json` and is
   reset whenever the batch token changes (i.e., a new manual run or a
   new watcher scan-cycle).
3. The `discount_from` and `discount_until` values are no longer read
   from defaults or user input. They are auto-derived per-file from the
   XLSX header row (datetime objects starting at column 8). Any value
   present in defaults for these two keys is overridden on a per-file
   basis right after the workbook is loaded and before the DOCX is filled.
4. `team3_email` and `team4_email` are now rendered as clickable
   `mailto:` hyperlinks (blue, underlined) in the generated DOCX/PDF.
   They were previously plain text. `team_email` is unchanged — it was
   already a hyperlink in the original template.

Existing behaviour preserved
----------------------------
- Watch mode, PDF export, emailer integration, XLSX table appending,
  template-fill mechanics, and all other placeholders are untouched.

Usage:
    python smartpayer_letter_generator.py --xlsx <path/to/file.xlsx>
    python smartpayer_letter_generator.py --watch <folder>
    python smartpayer_letter_generator.py --setup-defaults
    python smartpayer_letter_generator.py --show-defaults
"""

import os
import sys
import time
import io
import json
import re
import uuid
import zipfile
import argparse
import calendar
import subprocess
import tempfile
import threading
from pathlib import Path
from datetime import datetime, date
from copy import deepcopy
from contextlib import contextmanager
from calendar import monthrange

# Import msvcrt for windows and fcntl for Linux/MacOS
if sys.platform == "win32":
    import msvcrt
else: 
    # Import fcntl for Linux/MacOS
    import fcntl

# Force UTF-8 output on Windows so Unicode characters in log lines don't crash
if sys.stdout and hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
if sys.stderr and hasattr(sys.stderr, "reconfigure"):
    try:
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

# Safe print that never raises UnicodeEncodeError on Windows cp1252 consoles
def _print(*args, **kwargs):
    try:
        print(*args, **kwargs)
    except UnicodeEncodeError:
        safe = " ".join(str(a).encode("ascii", "replace").decode("ascii") for a in args)
        print(safe, **{k: v for k, v in kwargs.items() if k != "end"})


# ─── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR    = Path(__file__).parent.resolve()
DEFAULTS_FILE = SCRIPT_DIR / "smartpayer_letter_defaults.json"
TEMPLATE_FILE = SCRIPT_DIR / "Smart_Payer_Program_Letter_Template.docx"
OUTPUT_DIR    = SCRIPT_DIR / "Generated_Letters"

# State file used to implement the per-batch `nnn` counter for letter_number.
# Format: {"batch_token": "<uuid>", "date": "YYYY-MM-DD", "counter": <int>}
BATCH_STATE_FILE = SCRIPT_DIR / ".smartpayer_batch_state.json"

_batch_counter_lock = threading.Lock()

# ─── Variable schema ───────────────────────────────────────────────────────────
# NOTE: `letter_date_num` has been REPLACED by `year_num` (yy of current year).
# `discount_from` and `discount_until` are kept here only so they appear in
# defaults/CLI flows, but in run_pipeline() they are ALWAYS overridden with
# values read from the input XLSX. They are not editable in the GUI anymore.
VARIABLE_SCHEMA = [
    {"key": "company_name",         "label": "Company Full Name",                        "example": "Mencari Cinta Sejati"},
    {"key": "company_abbr",         "label": "Company Abbreviation",                     "example": "MCS"},
    {"key": "letter_date",          "label": "Letter Date (dd Month yyyy)",              "example": "26 September 2026"},
    {"key": "month_name",           "label": "Month Name (auto from filename)",         "example": "April"},
    {"key": "month_year",           "label": "Month & Year (auto from filename)",       "example": "April 2026"},
    {"key": "month_roman",          "label": "Month (Roman Numerals)",                   "example": "IX"},
    {"key": "letter_number",        "label": "Letter Number (mmddnnn)",                  "example": "0926001"},
    {"key": "year_num",             "label": "Last two digits of the year (yy)",         "example": "26"},
    {"key": "deposit_rate",         "label": "Deposit Interest Rate (%/year)",           "example": "2"},
    {"key": "deposit_days",         "label": "Deposit Duration (days)",                  "example": "3"},
    {"key": "daily_rate",           "label": "Daily Reward Rate (%/day)",                "example": "0.022"},
    {"key": "annual_rate",          "label": "Annualized Reward Rate (%/year)",          "example": "8"},
    {"key": "rate_multiplier_num",  "label": "Rate Multiplier (number)",                 "example": "4"},
    {"key": "rate_multiplier_word", "label": "Rate Multiplier (words/Indonesian)",       "example": "empat"},
    {"key": "end_of_month_date",    "label": "End of Next Month (full date, auto from filename)", "example": "31 May 2026"},
    {"key": "discount_from",        "label": "Discount Period From (dd)  [auto from XLSX]",     "example": "(auto-derived)"},
    {"key": "discount_until",       "label": "Discount Period Until (dd Month yyyy) [auto from XLSX]", "example": "(auto-derived)"},
    {"key": "team1_honorific_name", "label": "Finance Team 1 (Honorific + Name)",       "example": "Bapak John"},
    {"key": "team2_honorific_name", "label": "Finance Team 2 (Honorific + Name)",       "example": "Bapak Arthur"},
    {"key": "team_email",           "label": "Finance Team Email",                       "example": "dagwon@amq.com"},
    {"key": "team3_honorific_name", "label": "Finance Team 3 (Honorific + Name)",       "example": "Bapak Edward"},
    {"key": "team3_email",          "label": "Finance Team 3 Email",                     "example": "edward@amq.com"},
    {"key": "team4_honorific_name", "label": "Finance Team 4 (Honorific + Name)",       "example": "Ibu Amber"},
    {"key": "team4_email",          "label": "Finance Team 4 Email",                     "example": "amber@amq.com"},
]

# ─── Placeholder → variable key map ────────────────────────────────────────────
# The "Date Number (i.e., 26)" placeholder in the template originally meant
# the day of the month, but is now used for the LAST TWO DIGITS OF THE YEAR
# (per the new spec — the Letter No. now ends with /V/26 rather than /V/12).
# The DOCX template does NOT need to be edited because the placeholder string
# we search for is identical; only its semantic mapping changes here.
REPLACEMENTS = [
    ("dd Month Name yyyy (i.e., 26 September 2026)",                          "letter_date"),
    ("Month Name yyyy (i.e., September 2026)",                                "month_year"),
    ("Month Name (i.e., September)",                                          "month_name"),
    ("Month in Roman Number (i.e., IX)",                                      "month_roman"),
    ("Date Number (i.e., 26)",                                                "year_num"),   # ← REMAPPED: was letter_date_num
    ("mmddnnn (i.e., 0926123)",                                               "letter_number"),
    ("company name i.e., Mencari Cinta Sejati (MCS)",                        "company_name_full"),
    ("abbreviation of the company name i.e., MCS",                           "company_abbr"),
    ("current deposit interest rate (i.e., 2)",                              "deposit_rate"),
    ("amount of days (i.e., 3)",                                              "deposit_days"),
    ("interest rate per day (i.e., 0.022)",                                  "daily_rate"),
    ("annualized interest rate (i.e., 8)",                                   "annual_rate"),
    ("amount of x times of deposit interest rate in number (i.e., 4)",      "rate_multiplier_num"),
    ("amount of x times of deposit interest rate but spelt (i.e., empat)",  "rate_multiplier_word"),
    # NOTE: the template's actual text reads "...depending of the month +1 )"
    # (with a typo and stray "+1"). After our run-merge step it collapses to
    # "...depending of the month+1)". The string below matches that exact
    # post-merge form so the placeholder is actually substituted. This was a
    # latent bug in the previous code — the old pattern never matched and
    # the placeholder text leaked through unchanged. Required for change #3
    # (end_of_month_date now reflects the last day of NEXT month).
    ("end of month date and year yyyy (i.e., 30/31 depending of the month+1 2026)", "end_of_month_date"),
    ("offered discount date FROM (dd) i.e., 24",                             "discount_from"),
    ("offered discount date UNTIL (dd Month Name yyyy) i.e., 27 September 2026", "discount_until"),
    ("Honorifics + Name No. 1 i.e., Bapak John",                            "team1_honorific_name"),
    ("Honorifics + Name Team No.2 i.e., Bapak Arthur",                      "team2_honorific_name"),
    ("finance team email (i.e., dagwon@amq.com",                             "team_email"),
    ("Honorifics + Name Team No. 3 (i.e., Bapak Edward)",                   "team3_honorific_name"),
    ("team no.3 email i.e., edward@amq.com",                                 "team3_email"),
    ("Honorifics + Name Team No. 4 (i.e., Ibu Amber)",                      "team4_honorific_name"),
    ("Team no.4 email (i.e., amber@amq.com)",                                "team4_email"),
]


# =============================================================================
# DEFAULT MANAGEMENT
# =============================================================================

def load_defaults() -> dict:
    if DEFAULTS_FILE.exists():
        with open(DEFAULTS_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_defaults(data: dict):
    with open(DEFAULTS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    _print("\n[OK] Defaults saved to", DEFAULTS_FILE)


def show_defaults():
    defaults = load_defaults()
    if not defaults:
        _print("\n[!] No defaults set yet. Run with --setup-defaults to create them.")
        return
    _print("\n--- Current Default Variables ---")
    for v in VARIABLE_SCHEMA:
        val = defaults.get(v["key"], "(not set)")
        _print(f"  {v['label']:<52} -> {val}")
    _print("-" * 68)


def prompt_all_variables(prefill: dict = None) -> dict:
    _print("\n--- Fill Template Variables ---")
    _print("  Press ENTER to keep the shown default/example value.\n")
    values = {}
    prefill = prefill or {}
    for v in VARIABLE_SCHEMA:
        current = prefill.get(v["key"]) or v["example"]
        hint = f" [{current}]" if current else ""
        try:
            raw = input(f"  {v['label']}{hint}: ").strip()
        except (EOFError, KeyboardInterrupt):
            raw = ""
        values[v["key"]] = raw if raw else current
    return values


def setup_defaults():
    _print("\n=== Smart Payer - Default Variable Setup ===")
    existing = load_defaults()
    if existing:
        _print("  Existing defaults found. Press ENTER to keep each value.\n")
    values = prompt_all_variables(prefill=existing)
    save_defaults(values)


# =============================================================================
# AUTO-DATE COMPUTATION  (used when GUI's "Always use today's date" is on)
# =============================================================================

# Roman numeral lookup for months 1..12 — only ever 12 values so a literal
# tuple is clearer (and faster) than a generic int→roman converter.
_MONTH_ROMAN = ("I", "II", "III", "IV", "V", "VI", "VII", "VIII",
                "IX", "X", "XI", "XII")


def compute_today_fields(today: datetime | None = None) -> dict:
    """
    Return the date-derived fields that should overwrite user/default values
    when the "Always use today's date" checkbox is on.

    Note: the `nnn` suffix of `letter_number` is intentionally NOT set here.
    It is appended later inside run_pipeline() using the batch counter, so
    that the GUI can preview the prefix (mmdd) without committing a counter.

    Note: `month_name`, `month_year`, and `end_of_month_date` used to be
    derived from today, but they are now derived from the XLSX FILENAME via
    parse_filename_month() / compute_filename_fields(). They're still
    returned here as fallbacks for callers that don't have an XLSX path
    (e.g. the GUI's "preview today's values" label, or filenames that
    don't match the expected pattern).
    """
    today = today or datetime.now()

    # Last day of the month FOLLOWING the current month — kept as a
    # fallback only; the real end_of_month_date comes from the filename's
    # month, not today's.
    if today.month == 12:
        next_year, next_month = today.year + 1, 1
    else:
        next_year, next_month = today.year, today.month + 1
    _, last_day_of_next_month = calendar.monthrange(next_year, next_month)
    eom_next_full = datetime(next_year, next_month, last_day_of_next_month).strftime("%d %B %Y")

    return {
        "letter_date":       today.strftime("%d %B %Y"),       # 12 May 2026
        "month_name":        today.strftime("%B"),             # May  (fallback)
        "month_year":        today.strftime("%B %Y"),          # May 2026  (fallback)
        "month_roman":       _MONTH_ROMAN[today.month - 1],    # V
        "year_num":          today.strftime("%y"),             # 26
        # `letter_number_prefix` is the mmdd part only; the final field
        # `letter_number` is assembled as prefix + nnn after the batch
        # counter is consumed.
        "letter_number_prefix": today.strftime("%m%d"),         # 0512
        "end_of_month_date": eom_next_full,                     # 30 June 2026  (fallback)
    }


# Filename month/year parser.
#
# SmartPayer XLSXes follow the pattern:
#     "Smartpayer <Month Name> <YYYY> <Client>.xlsx"
# e.g. "Smartpayer April 2026 BANLY THEO.xlsx".
#
# The leading "Smartpayer" is optional in the regex so files renamed by
# the user (e.g. "April 2026 ACME.xlsx") still parse. The match is
# case-insensitive on the month name to tolerate "april" / "APRIL".
#
# We accept Indonesian month names too, since the user base is Indonesian.
# All twelve months are listed both in English and Bahasa.

_MONTH_NAMES_EN = ("January", "February", "March", "April", "May", "June",
                   "July", "August", "September", "October", "November", "December")
_MONTH_NAMES_ID = ("Januari", "Februari", "Maret", "April", "Mei", "Juni",
                   "Juli", "Agustus", "September", "Oktober", "November", "Desember")

# Build a case-insensitive lookup map: name → 1-based month index.
_MONTH_LOOKUP = {}
for _idx, _n in enumerate(_MONTH_NAMES_EN, start=1):
    _MONTH_LOOKUP[_n.lower()] = _idx
for _idx, _n in enumerate(_MONTH_NAMES_ID, start=1):
    _MONTH_LOOKUP[_n.lower()] = _idx

# Compile the month-name alternation once.
_FILENAME_MONTH_RE = __import__("re").compile(
    r"\b(?P<month>(?:" +
    "|".join(_MONTH_NAMES_EN + _MONTH_NAMES_ID) +
    r"))\s+(?P<year>\d{4})\b",
    flags=__import__("re").IGNORECASE
)


def parse_filename_month(filename: str) -> tuple[int, int] | None:
    """
    Extract (month_index, year) from an XLSX filename like
    "Smartpayer April 2026 BANLY THEO.xlsx".

    Returns (month, year) as (1..12, full year int), or None if no
    month-name + 4-digit-year pair appears in the filename.

    Tolerant of:
      - English or Indonesian month names
      - Any casing of the month name
      - Files without the leading "Smartpayer" prefix
      - Additional client-name junk after the year
    """
    if not filename:
        return None
    m = _FILENAME_MONTH_RE.search(str(filename))
    if not m:
        return None
    month_idx = _MONTH_LOOKUP.get(m.group("month").lower())
    if month_idx is None:
        return None
    try:
        year = int(m.group("year"))
    except ValueError:
        return None
    if not (1900 <= year <= 2100):  # sanity bound
        return None
    return month_idx, year


def compute_filename_fields(filename: str) -> dict:
    """
    Derive {month_name, month_year, end_of_month_date} from an XLSX
    filename's encoded billing month.

    For "Smartpayer April 2026 ...":
        month_name        = "April"
        month_year        = "April 2026"
        end_of_month_date = "31 May 2026"           ← last day of next month,
                                                       full "<dd> <Mon> <YYYY>"

    Returns {} if the filename doesn't carry a month/year — caller should
    fall back to compute_today_fields() values in that case.
    """
    parsed = parse_filename_month(filename)
    if parsed is None:
        return {}
    month_idx, year = parsed

    # Use English month names in the output (template body is in Indonesian
    # but uses English month names like "May 2026" in the sample defaults).
    month_name = _MONTH_NAMES_EN[month_idx - 1]
    month_year = f"{month_name} {year}"

    # NEXT month: roll over December → January of next year.
    if month_idx == 12:
        next_month_idx, next_year = 1, year + 1
    else:
        next_month_idx, next_year = month_idx + 1, year
    _, last_day = calendar.monthrange(next_year, next_month_idx)
    next_month_name = _MONTH_NAMES_EN[next_month_idx - 1]
    end_of_month_date = f"{last_day} {next_month_name} {next_year}"

    return {
        "month_name":        month_name,
        "month_year":        month_year,
        "end_of_month_date": end_of_month_date,
    }


# =============================================================================
# PER-BATCH COUNTER FOR letter_number
# =============================================================================
#
# Design
# ------
# - The GUI / watcher generates one fresh `batch_token` (uuid4) per batch:
#     * "Manual Generate Letter" click → one token, one file → counter = 001.
#     * Watcher scan-cycle that picks up N files → one shared token, files
#       1..N → counters 001..N (incremented serially since the watcher
#       blocks on proc.wait() per file).
# - Each subprocess invocation of this script receives `--batch-token <uuid>`.
# - On invocation we read .smartpayer_batch_state.json:
#     * If the stored token matches AND the stored date == today: increment
#       the counter and persist.
#     * Otherwise: reset counter to 1 and persist with the new token/date.
# - The counter is therefore guaranteed to start at 001 for each new batch
#   and to NEVER carry over across days even within the same batch token
#   (cheap defence-in-depth against an extremely long-running watcher).


def read_batch_state() -> dict:
    """Read batch state file. Callers should hold _batch_counter_lock."""
    if not BATCH_STATE_FILE.exists():
        return {}
    try:
        with open(BATCH_STATE_FILE, "r", encoding="utf-8") as f:
            content = f.read().strip()
            if not content:
                return {}
            return json.loads(content)
    except (json.JSONDecodeError, UnicodeDecodeError, OSError):
        _print(f"  [!] Batch state file corrupted - resetting", "warning")
        return {}

def write_batch_state(state: dict):
    """Write batch state file. Callers should hold _batch_counter_lock."""
    BATCH_STATE_FILE.parent.mkdir(parents=True, exist_ok=True)
    try:
        # Write directly, overwriting previous content
        with open(BATCH_STATE_FILE, "w", encoding="utf-8") as f:
            json.dump(state, f, indent=2)
            f.flush()
            os.fsync(f.fileno())
    except Exception as e:
        _print(f"  [!] Failed to write batch state: {e}", "error")
        raise

def consume_batch_counter(batch_token: str | None) -> int:
    """
    Atomically increment-or-reset the per-batch counter for `batch_token`.
    
    Supports two token formats:
      - "daily-YYYYMMDD" → counter persists all day, resets at midnight
      - UUID (e.g., "a1b2c3d4...") → counter resets when token changes
    
    Thread-safe via _batch_counter_lock (for same-process threads).
    """
    if not batch_token:
        return 1

    with _batch_counter_lock:
        today_str = datetime.now().strftime("%Y-%m-%d")
        today_date_token = f"daily-{datetime.now().strftime('%Y%m%d')}"
        state = read_batch_state()

        # Determine if we should increment or reset
        stored_token = state.get("batch_token")
        stored_date = state.get("date")
        
        if batch_token.startswith("daily-"):
            # Daily mode: reset only when date changes
            expected_date = batch_token.replace("daily-", "")
            current_date = datetime.now().strftime("%Y%m%d")
            if stored_token == batch_token and expected_date == current_date:
                counter = int(state.get("counter", 0)) + 1
            else:
                counter = 1
        else:
            # UUID mode: reset when token or date changes
            if stored_token == batch_token and stored_date == today_str:
                counter = int(state.get("counter", 0)) + 1
            else:
                counter = 1

        new_state = {
            "batch_token": batch_token,
            "date": today_str,
            "counter": counter,
        }
        write_batch_state(new_state)
        
        _print(f"  [✓] Letter sequence: {counter:03d} (batch: {batch_token[:12]}…)", "info")
        return counter

def reset_batch_counter(batch_token: str):
    """Reset batch counter. Thread-safe via _batch_counter_lock."""
    with _batch_counter_lock:
        state = {
            "batch_token": batch_token,
            "date": datetime.now().strftime("%Y-%m-%d"),
            "counter": 0,
        }
        write_batch_state(state)
    _print(f"  [OK] Batch counter reset for token {batch_token[:8]}…", "success")

# =============================================================================
# REWARD-RATE DERIVATION  (from XLSX daily_rate + deposit_rate)
# =============================================================================
#
# When the input XLSX provides the daily reward rate as a float in its
# top-right header cell (e.g. 0.00022 above the "Simulasi Imbalan" label),
# the following four template fields can all be derived from it and the
# user's configured deposit_rate, instead of being hand-edited.
#
#   daily_rate           = float-from-XLSX × 100         (e.g. 0.022   %/day)
#   annual_rate          = daily_rate × 365              (e.g. 8.03    %/year)
#   rate_multiplier_num  = round(annual_rate / deposit_rate)   (e.g. 4)
#   rate_multiplier_word = int_to_indonesian(...)        (e.g. "empat")
#
# This block adds:
#   * int_to_indonesian()  — small helper for 0..9999 (covers all realistic
#     deposit-rate multipliers; expandable to millions if ever needed)
#   * compute_rate_fields() — returns the four string-typed template fields
#     ready to drop into `variables`.

# Indonesian numerals for integer spelling. Standard Bahasa convention:
#   - "se-" prefix for one in the tens/hundreds/thousands positions
#     (sepuluh, sebelas, seratus, seribu)
#   - 11..19 use "belas" suffix
_ID_DIGITS = ("nol", "satu", "dua", "tiga", "empat",
              "lima", "enam", "tujuh", "delapan", "sembilan")
_ID_TEENS  = ("sepuluh", "sebelas", "dua belas", "tiga belas", "empat belas",
              "lima belas", "enam belas", "tujuh belas", "delapan belas",
              "sembilan belas")


def int_to_indonesian(n: int) -> str:
    """
    Spell a non-negative integer 0..9999 in Bahasa Indonesia.

    Examples:
        0    -> "nol"
        1    -> "satu"
        4    -> "empat"
        11   -> "sebelas"
        21   -> "dua puluh satu"
        100  -> "seratus"
        345  -> "tiga ratus empat puluh lima"
        1000 -> "seribu"
        2024 -> "dua ribu dua puluh empat"
    """
    if n < 0:
        return f"minus {int_to_indonesian(-n)}"
    if n < 10:
        return _ID_DIGITS[n]
    if n < 20:
        return _ID_TEENS[n - 10]
    if n < 100:
        tens, ones = divmod(n, 10)
        head = _ID_DIGITS[tens] + " puluh"
        return head if ones == 0 else f"{head} {_ID_DIGITS[ones]}"
    if n < 1000:
        hundreds, rest = divmod(n, 100)
        head = "seratus" if hundreds == 1 else f"{_ID_DIGITS[hundreds]} ratus"
        return head if rest == 0 else f"{head} {int_to_indonesian(rest)}"
    if n < 10000:
        thousands, rest = divmod(n, 1000)
        head = "seribu" if thousands == 1 else f"{_ID_DIGITS[thousands]} ribu"
        return head if rest == 0 else f"{head} {int_to_indonesian(rest)}"
    # Above 9999: fall back to a recursive ribu/juta build. Realistically we
    # never expect rate multipliers above a handful, so this branch is just
    # defensive — keeps the function total.
    if n < 1_000_000:
        thousands, rest = divmod(n, 1000)
        head = f"{int_to_indonesian(thousands)} ribu"
        return head if rest == 0 else f"{head} {int_to_indonesian(rest)}"
    millions, rest = divmod(n, 1_000_000)
    head = f"{int_to_indonesian(millions)} juta"
    return head if rest == 0 else f"{head} {int_to_indonesian(rest)}"


def _fmt_rate_number(x: float) -> str:
    """
    Render a percentage-style float as a clean string:
      - whole numbers come out without a decimal point   (8.0 -> "8")
      - otherwise up to 3 decimals, trailing zeros stripped
        (0.022 -> "0.022", 0.0220 -> "0.022", 8.03 -> "8.03")
    """
    if x == int(x):
        return str(int(x))
    s = f"{x:.3f}".rstrip("0").rstrip(".")
    return s or "0"


def clean_template_value(value) -> str:
    """
    Render template replacement values without cosmetic trailing decimals.
    Keeps meaningful decimal precision such as 0.022, but turns 8.0 / 2.00
    into 8 / 2 before writing the DOCX XML.
    """
    if value is None:
        return ""
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:.6f}".rstrip("0").rstrip(".")
    text = str(value)
    if re.fullmatch(r"-?\d+\.0+", text):
        return text.split(".", 1)[0]
    if re.fullmatch(r"-?\d+\.\d+", text):
        return text.rstrip("0").rstrip(".")
    return text


def compute_rate_fields(daily_rate_decimal: float,
                        deposit_rate: str | float | int | None) -> dict:
    """
    Derive the four reward-rate template fields from the XLSX daily rate
    (a decimal like 0.00022) and the user-configured deposit_rate (a string
    like "2" representing 2 %/year, kept as a string elsewhere in this code
    to match the rest of `variables`).

    Returns a dict with string-typed values so it can be `.update()`'d
    directly into `variables`. Returns {} if daily_rate_decimal is missing
    or non-numeric — the caller should then fall back to defaults.

    The annual rate is rounded to a whole number when the result is
    sufficiently close (within 0.05 percentage points) to avoid printing
    "8.03 %/year" for a multiplier that's clearly meant to be exactly 4×.
    Otherwise we keep up to 2 decimals.
    """
    try:
        decimal = float(daily_rate_decimal)
    except (TypeError, ValueError):
        return {}
    if decimal <= 0:
        return {}

    daily_pct  = decimal * 100              # 0.00022 -> 0.022
    annual_pct = daily_pct * 365            # 0.022 * 365 = 8.03

    # Snap near-integer annual rates to whole numbers (within 0.1).
    annual_rounded = round(annual_pct)
    if abs(annual_pct - annual_rounded) < 0.1:
        annual_pct_out = float(annual_rounded)
    else:
        annual_pct_out = round(annual_pct, 2)

    try:
        deposit_val = float(str(deposit_rate).strip()) if deposit_rate not in (None, "") else 0.0
    except (TypeError, ValueError):
        deposit_val = 0.0

    if deposit_val > 0:
        multiplier = round(annual_pct_out / deposit_val)
    else:
        multiplier = 0

    return {
        "daily_rate":           _fmt_rate_number(daily_pct),
        "annual_rate":          _fmt_rate_number(annual_pct_out),
        "rate_multiplier_num":  str(multiplier),
        "rate_multiplier_word": int_to_indonesian(multiplier) if multiplier >= 0 else "nol",
    }


# =============================================================================
# XML RUN MERGER  (lxml-based — safe, no external scripts needed)
# =============================================================================

W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WQ = f"{{{W}}}"  # Clark-notation prefix


def _etree_to_string(elem) -> bytes:
    """Serialise an lxml element back to bytes, preserving the declaration."""
    from lxml import etree
    return etree.tostring(elem, xml_declaration=False, encoding="unicode").encode("utf-8")


def _rpr_key(run_el) -> str:
    """Return a canonical string key for the rPr of a run, for comparison."""
    from lxml import etree
    rpr = run_el.find(f"{WQ}rPr")
    if rpr is None:
        return ""
    return etree.tostring(rpr, encoding="unicode")


def _run_has_special(run_el) -> bool:
    """Return True if the run contains elements we must not merge across."""
    skip_tags = {f"{WQ}{t}" for t in (
        "br", "instrText", "fldChar", "rStyle", "lastRenderedPageBreak",
        "sym", "pgNum", "tab",
    )}
    return any(child.tag in skip_tags for child in run_el)


def _merge_paragraph_runs(para_el):
    """
    In-place merge of adjacent <w:r> siblings inside a paragraph element
    that share the same <w:rPr>. Skips over <w:proofErr> and <w:bookmarkStart/End>
    elements — these are safe to remove when merging surrounding runs.
    Accumulates <w:t> text across compatible runs.
    """
    from lxml import etree

    SKIP_BETWEEN = {f"{WQ}{t}" for t in ("proofErr", "bookmarkStart", "bookmarkEnd")}

    def children_list():
        return list(para_el)

    changed = True
    while changed:
        changed = False
        children = children_list()
        i = 0
        while i < len(children):
            child = children[i]
            if child.tag != f"{WQ}r" or _run_has_special(child):
                i += 1
                continue

            key = _rpr_key(child)

            # Look ahead past skippable elements for a compatible run
            j = i + 1
            skip_els = []
            while j < len(children):
                nxt = children[j]
                if nxt.tag in SKIP_BETWEEN:
                    skip_els.append(nxt)
                    j += 1
                    continue
                if nxt.tag == f"{WQ}r" and not _run_has_special(nxt) and _rpr_key(nxt) == key:
                    # Merge: absorb nxt text into child
                    for t in nxt.findall(f"{WQ}t"):
                        existing_t = child.find(f"{WQ}t")
                        if existing_t is None:
                            new_t = etree.SubElement(child, f"{WQ}t")
                            new_t.text = t.text or ""
                        else:
                            existing_t.text = (existing_t.text or "") + (t.text or "")
                    for sk in skip_els:
                        para_el.remove(sk)
                    para_el.remove(nxt)
                    changed = True
                    skip_els = []
                    children = children_list()
                    break
                else:
                    break

            # Fix xml:space on the surviving <w:t>
            t_el = child.find(f"{WQ}t")
            if t_el is not None:
                txt = t_el.text or ""
                space_ns = "http://www.w3.org/XML/1998/namespace"
                if txt != txt.strip() or "  " in txt:
                    t_el.set(f"{{{space_ns}}}space", "preserve")
                else:
                    t_el.attrib.pop(f"{{{space_ns}}}space", None)

            i += 1


def merge_all_runs(xml_bytes: bytes) -> bytes:
    """
    Parse the document XML with lxml, merge runs in every paragraph,
    and return the serialised bytes.
    """
    from lxml import etree
    root = etree.fromstring(xml_bytes)
    body = root.find(".//{%s}body" % W)
    if body is None:
        return xml_bytes
    for para in root.iter(f"{WQ}p"):
        _merge_paragraph_runs(para)
    return etree.tostring(root, xml_declaration=True,
                          encoding="UTF-8", standalone=True)


# =============================================================================
# DOCX TEMPLATE FILLING
# =============================================================================

def fill_template(variables: dict, template_path: Path, output_path: Path):
    """
    1. Read the DOCX zip
    2. Merge split runs in document.xml (no external script needed)
    3. Do text replacements for all variables
    4. Strip yellow highlights
    5. Remove the blank Lampiran page from the original template
    6. Repack as a new DOCX
    """
    variables = {k: clean_template_value(v) for k, v in dict(variables).items()}
    company_name = variables.get("company_name", "")
    company_abbr = variables.get("company_abbr", "")
    variables["company_name_full"] = (
        f"{company_name} ({company_abbr})" if company_name and company_abbr
        else company_name or company_abbr
    )

    # Read original zip into memory
    with zipfile.ZipFile(template_path, "r") as zin:
        names = zin.namelist()
        file_bytes = {name: zin.read(name) for name in names}

    # Step 1: merge split runs with lxml so placeholders become contiguous
    merged_bytes = merge_all_runs(file_bytes["word/document.xml"])

    # Steps 2-5 work on the string form for simple text substitution
    xml = merged_bytes.decode("utf-8")

    # Step 2: variable replacements
    for placeholder, var_key in REPLACEMENTS:
        replacement = variables.get(var_key, "")
        if replacement:
            xml = xml.replace(placeholder, replacement)

    # Step 2b: collapse the redundant `<end_of_month_date> <month_year>` pair.
    #
    # The template body says "... sampai dengan {end of month date} {Month
    # Name yyyy}, untuk dibayarkan ..." — two placeholders side by side.
    # Historically `end_of_month_date` was a single day-of-month integer
    # (e.g. "30") and `month_year` immediately after it supplied the month
    # name + year (e.g. "May 2026") to read as "... sampai dengan 30 May
    # 2026, ...".
    #
    # Now `end_of_month_date` is itself a full date string (e.g. "31 May
    # 2026") derived from the XLSX filename's next-month rollover, and
    # `month_year` is the filename's billing month (e.g. "April 2026").
    # If we leave both placeholders in place the sentence would read
    # "... sampai dengan 31 May 2026 April 2026, ..." — wrong.
    #
    # The two values often land in SEPARATE <w:r> runs separated by a
    # whitespace-only run, so a plain regex on the text won't match. We
    # do the collapse at the XML level: find every `>{eom}<` end-tag
    # boundary, then look ahead for an immediate `>{month_year}<` boundary
    # within a short window of intervening XML tags. When we find that
    # adjacency, we blank out the contents of the second <w:t> element
    # (so the redundant phrase disappears) without disturbing the runs'
    # formatting or any surrounding bookmarks.
    eom = variables.get("end_of_month_date", "").strip()
    my  = variables.get("month_year", "").strip()
    if eom and my and eom != my:
        # Match: ">{eom}</w:t>"  ...short window of XML tags + whitespace runs...
        #        "<w:t...>{my}</w:t>"
        # We DON'T touch the eom side; we blank out the my-side <w:t>'s text
        # AND drop any whitespace-only <w:t> runs sitting between them so we
        # don't end up with "31 May 2026 ," (note the leading space before
        # the comma) in the visible text. We use the verbose `_drop_redundant`
        # callback rather than a single regex substitution because we need
        # to walk through and stitch the cleaned-up middle back together.
        collapse_pattern = re.compile(
            r">" + re.escape(eom) + r"</w:t>"
            # up to ~400 chars of intervening tags + whitespace-only text
            r"(?P<between>(?:[^<]*<[^>]*>){0,30}?)"
            r"<w:t(?P<attrs>[^>]*)>" + re.escape(my) + r"</w:t>",
            flags=re.DOTALL
        )

        # Sub-pattern: a <w:r> whose only <w:t> contains pure whitespace.
        # Used to strip those out of the "between" capture.
        whitespace_run = re.compile(
            r"<w:r\b[^>]*>(?:\s*<w:rPr>.*?</w:rPr>)?\s*"
            r"<w:t[^>]*>\s+</w:t>\s*</w:r>",
            flags=re.DOTALL
        )

        def _drop_redundant_my(m):
            attrs = m.group("attrs") or ""
            between = m.group("between")
            between_clean = whitespace_run.sub("", between)
            return (">" + eom + "</w:t>"
                    + between_clean
                    + "<w:t" + attrs + "></w:t>")

        new_xml, n_subs = collapse_pattern.subn(_drop_redundant_my, xml)
        if n_subs > 0:
            xml = new_xml

    # Step 3: email hyperlink — the email itself sits in a <w:t> inside <w:hyperlink>
    team_email = variables.get("team_email", "")
    if team_email:
        xml = xml.replace(">dagwon@amq.com<", f">{team_email}<")
        xml = xml.replace("finance team email (i.e., ", "")

    # Step 4: strip all highlights
    xml = re.sub(r'\s*<w:highlight[^/]*/>', "", xml)

    # Step 5: remove the explicit page-break paragraph + standalone Lampiran heading
    # that creates the blank page 3. Our append step adds its own page break + title.
    xml = re.sub(
        r'<w:p\b[^>]*>\s*<w:r>\s*<w:br\s+w:type="page"/>\s*</w:r>\s*</w:p>',
        "", xml
    )
    xml = re.sub(
        r'<w:p\b[^>]*>(?:\s*<w:(?:pPr|bookmarkStart|bookmarkEnd)[^>]*>.*?</w:[^>]+>)*'
        r'\s*<w:r\b[^>]*>(?:\s*<w:rPr>.*?</w:rPr>)?\s*(?:<w:lastRenderedPageBreak/>\s*)?'
        r'<w:t[^>]*>Lampiran Penghitungan Imbalan Smart Payer</w:t>\s*</w:r>'
        r'(?:\s*<w:(?:bookmarkStart|bookmarkEnd)[^>]*/?>)*\s*</w:p>',
        "", xml, flags=re.DOTALL
    )
    xml = re.sub(r'<w:p\b[^>]*w14:paraId="2E3EAA17"[^>]*/?>',  "", xml)

    file_bytes["word/document.xml"] = xml.encode("utf-8")

    # Write new zip
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, file_bytes[name])

    _print(f"  [OK] Filled DOCX: {output_path.name}")


# =============================================================================
# XLSX TABLE EXTRACTION
# =============================================================================

def extract_xlsx_data(xlsx_path: Path) -> dict:
    """
    Parse the SmartPayer XLSX.

    Column layout (0-indexed):
      0  Payer
      1  Bill To Party
      2  Payer Name
      3  Document No
      4  Doc. Date
      5  Net Due Date
      6  URL-1v.Ostd. (also seen as "Tot.Inv.Ostd.")
      7  DPP
      8+ Date columns (payment simulation dates)
      last non-None header col before dates: daily rate
    """
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb.active

    rows_data, date_columns = [], []
    payer_name = daily_rate = None
    pph_row = net_disc_row = transfer_row = {}

    all_rows = list(ws.iter_rows(values_only=True))
    if not all_rows:
        return {}

    header = all_rows[0]

    # Find date columns: header cols that are datetime objects (start from col 8)
    # Also remember the raw datetime objects (not just the formatted strings) so
    # callers can derive discount_from / discount_until from them later.
    date_col_indices = []
    date_objects = []
    for ci in range(8, len(header)):
        val = header[ci]
        if isinstance(val, (datetime, date)):
            date_columns.append(val.strftime("%d %b %Y"))
            date_col_indices.append(ci)
            date_objects.append(val)
        elif val is not None and not isinstance(val, str):
            # Could be a numeric date serial — treat as daily rate placeholder
            pass

    # Daily rate: last column header that is a small float (e.g. 0.00022)
    for ci in range(len(header) - 1, 7, -1):
        v = header[ci]
        if isinstance(v, float) and 0 < v < 1:
            daily_rate = v
            break

    sum_ostd = 0.0
    sum_dpp = 0.0
    sum_rewards = [0.0] * len(date_col_indices)

    for row in all_rows[1:]:
        label_g = str(row[6] or "")
        if "PPH" in label_g.upper() or "PPh" in label_g:
            pph_row = {"rewards": [row[ci] for ci in date_col_indices]}
            continue
        if "Net Disc" in label_g:
            net_disc_row = {"rewards": [row[ci] for ci in date_col_indices]}
            continue
        if "transfer" in label_g.lower() or "nominal" in label_g.lower():
            transfer_row = {"rewards": [row[ci] for ci in date_col_indices]}
            continue

        if row[3] is None:
            continue

        doc_date = row[4]
        due_date = row[5]
        dpp_val  = row[7]

        if row[2] is not None:
            payer_name = row[2]

        reward_vals = [row[ci] for ci in date_col_indices]

        rows_data.append({
            "payer_name": str(row[2] or ""),
            "doc_no":     row[3],
            "doc_date":   doc_date.strftime("%d/%m/%Y") if isinstance(doc_date, (datetime, date)) else str(doc_date or ""),
            "due_date":   due_date.strftime("%d/%m/%Y") if isinstance(due_date, (datetime, date)) else str(due_date or ""),
            "ostd":       row[6],
            "dpp":        dpp_val,
            "rewards":    reward_vals,
        })

        try:
            sum_ostd += float(row[6] or 0)
        except (TypeError, ValueError):
            pass
        try:
            sum_dpp += float(dpp_val or 0)
        except (TypeError, ValueError):
            pass
        for j, rv in enumerate(reward_vals):
            try:
                sum_rewards[j] += float(rv or 0)
            except (TypeError, ValueError):
                pass

    totals = {
        "ostd":    sum_ostd if sum_ostd else None,
        "dpp":     sum_dpp if sum_dpp else None,
        "rewards": [v if v else None for v in sum_rewards],
    }

    return {
        "payer_name":    payer_name,
        "date_columns":  date_columns,
        "date_col_indices": date_col_indices,
        "date_objects":  date_objects,    # NEW: raw datetime list for discount range derivation
        "daily_rate":    daily_rate,
        "rows":          rows_data,
        "totals":        totals,
        "pph_row":       pph_row,
        "net_disc_row":  net_disc_row,
        "transfer_row":  transfer_row,
    }


def derive_discount_range(xlsx_data: dict) -> tuple[str, str] | tuple[None, None]:
    """
    Inspect the parsed XLSX data and return (discount_from, discount_until)
    derived from the dated header columns.

      discount_from  = earliest_date.strftime("%d")            e.g. "27"
      discount_until = latest_date.strftime("%d %B %Y")        e.g. "30 April 2026"

    Returns (None, None) if no datetime columns were found, in which case the
    caller should fall back to whatever was in the defaults.
    """
    dates = sorted(xlsx_data.get("date_objects") or [])
    if not dates:
        return None, None
    return dates[0].strftime("%d"), dates[-1].strftime("%d %B %Y")


def fmt_num(val):
    if val is None:
        return "-"
    try:
        f = float(val)
        return f"{int(round(f)):,}"
    except Exception:
        return str(val)


# =============================================================================
# DOCX TABLE APPENDING
# =============================================================================

def _set_table_page_margins(doc):
    """
    Keep A4 portrait (21 x 29.7 cm) but shrink left/right margins to 1.0 cm
    so the table has ~19 cm usable width.  Top/bottom stay at 1.5 cm to keep
    the inherited letterhead header/footer visible.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    body   = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        sectPr = OxmlElement("w:sectPr")
        body.append(sectPr)

    pgSz = sectPr.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        sectPr.insert(0, pgSz)
    pgSz.set(qn("w:w"), "11906")
    pgSz.set(qn("w:h"), "16838")
    pgSz.attrib.pop(qn("w:orient"), None)

    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is None:
        pgMar = OxmlElement("w:pgMar")
        sectPr.append(pgMar)
    pgMar.set(qn("w:top"),    "851")
    pgMar.set(qn("w:bottom"), "851")
    pgMar.set(qn("w:left"),   "567")
    pgMar.set(qn("w:right"),  "567")
    pgMar.set(qn("w:header"), "709")
    pgMar.set(qn("w:footer"), "709")


# ── NEW: convert team3/team4 plain-text emails into mailto hyperlinks ────────
def _convert_emails_to_hyperlinks(doc, emails: list[str]):
    """
    Find every <w:r> in the document whose visible text is exactly one of
    `emails`, and wrap each such run inside a new <w:hyperlink> element
    pointing to mailto:<that email>. The run's character style is set to
    "Hyperlink" so Word renders it blue + underlined.

    We use python-docx's part.relate_to(...) to register the external mailto
    relationship and get back the rId used by the new w:hyperlink element.

    Why match the exact run text rather than substrings:
      - team3_email / team4_email values are substituted into a previously-
        plain-text run by fill_template(). That run contains *only* the
        email address by the time we get here, so a strict equality match
        is unambiguous and avoids accidentally hyperlinking the same
        address elsewhere (e.g. in body text).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    # De-dupe and drop falsy values defensively.
    targets = {e.strip() for e in emails if e and e.strip()}
    if not targets:
        return

    main_part = doc.part
    doc_root  = doc.element  # <w:document>

    # Cache mailto → rId so we don't create duplicate relationships if the
    # same email appears more than once in the document.
    rid_for_email: dict[str, str] = {}

    def _ensure_rid(email: str) -> str:
        if email in rid_for_email:
            return rid_for_email[email]
        rid = main_part.relate_to(
            f"mailto:{email}",
            RT.HYPERLINK,
            is_external=True,
        )
        rid_for_email[email] = rid
        return rid

    # Iterate every run in the body. We rebuild the list each iteration
    # because we are mutating the tree (replacing <w:r> with <w:hyperlink>).
    # Doing it in one pass is fine: we only replace runs that match, and
    # the wrapped run stays in the same position so the next outer-loop
    # iteration still sees the rest of the runs.
    runs = list(doc_root.iter(qn("w:r")))
    for r in runs:
        # Collect the visible text of this run by joining all child <w:t>.
        texts = r.findall(qn("w:t"))
        if not texts:
            continue
        visible = "".join((t.text or "") for t in texts).strip()
        if visible not in targets:
            continue

        email = visible
        rid   = _ensure_rid(email)

        # Apply the "Hyperlink" character style so Word renders blue/underline.
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r.insert(0, rPr)
        # Replace any existing rStyle to avoid duplicates.
        for old_rstyle in rPr.findall(qn("w:rStyle")):
            rPr.remove(old_rstyle)
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.insert(0, rStyle)

        # Build the wrapping hyperlink element.
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rid)
        hyperlink.set(qn("w:history"), "1")

        # Insert the hyperlink in place of the run and move the run inside it.
        parent = r.getparent()
        if parent is None:
            continue
        idx = list(parent).index(r)
        parent.insert(idx, hyperlink)
        hyperlink.append(r)


def append_tables_to_docx(docx_path: Path, xlsx_data: dict, output_path: Path,
                          hyperlink_emails: list[str] | None = None):
    """
    Open the filled DOCX, append the Lampiran tables, optionally convert
    `hyperlink_emails` into clickable mailto: hyperlinks, and save.

    Adding the email hyperlinking as a step in this function (rather than
    a separate doc-load round-trip) avoids reopening the document twice.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(docx_path)

    # ── Detect body font from existing paragraphs ─────────────────────────────
    body_font, body_size = "Calibri", 11
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:  body_font = run.font.name
            if run.font.size:  body_size = int(run.font.size.pt)
            break
        break

    # ── Cell helpers ──────────────────────────────────────────────────────────

    def set_cell_text(cell, text, bold=False, size=8, color=None,
                      italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = align
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run = para.add_run(str(text))
        run.font.name   = body_font
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)

    def shade_cell(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        tcPr.append(shd)

    def set_cell_margins(cell, top=40, bottom=40, left=50, right=50):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        mar  = tcPr.find(qn("w:tcMar"))
        if mar is None:
            mar = OxmlElement("w:tcMar")
            tcPr.append(mar)
        for side, val in [("top", top), ("bottom", bottom),
                          ("left", left), ("right", right)]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"),    str(val))
            el.set(qn("w:type"), "dxa")
            old = mar.find(qn(f"w:{side}"))
            if old is not None:
                mar.remove(old)
            mar.append(el)

    def write_merged_label(cell, text, bold=False, size=8,
                           align=WD_ALIGN_PARAGRAPH.CENTER, shade=None):
        """
        Write a label into a *merged* cell, replacing any leftover empty
        paragraphs that python-docx's merge() carried over from the cells
        we just merged in. Without this, each former cell contributes its
        own <w:p> to the merged cell and the row balloons to 5-6x the
        expected height. This is the root-cause fix for the bloated
        Total / Total PPH 23 / Net Discount / Nominal Transfer rows.

        We do this by directly editing the cell's XML: drop every existing
        <w:p> child of <w:tc>, then call set_cell_text() which adds a
        single fresh paragraph with the label.
        """
        tc = cell._tc
        for p in list(tc.findall(qn("w:p"))):
            tc.remove(p)
        # set_cell_text() will create one new paragraph via cell.paragraphs[0]
        # — but since we just removed every <w:p>, cell.paragraphs is empty.
        # python-docx auto-recreates one when we assign cell.text. Let it.
        set_cell_text(cell, text, bold=bold, size=size, align=align)
        if shade is not None:
            shade_cell(cell, shade)
        set_cell_margins(cell)

    def set_table_borders(table):
        tbl   = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "A0A0A0")
            tblBorders.append(el)
        old = tblPr.find(qn("w:tblBorders"))
        if old is not None:
            tblPr.remove(old)
        tblPr.append(tblBorders)

    def apply_col_widths_cm(table, cm_widths):
        CM = 567
        tbl   = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        total_twips = int(sum(cm_widths) * CM)
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"),    str(total_twips))
        tblW.set(qn("w:type"), "dxa")

        tblLayout = tblPr.find(qn("w:tblLayout"))
        if tblLayout is None:
            tblLayout = OxmlElement("w:tblLayout")
            tblPr.append(tblLayout)
        tblLayout.set(qn("w:type"), "fixed")

        old_grid = tbl.find(qn("w:tblGrid"))
        if old_grid is not None:
            tbl.remove(old_grid)
        tblGrid = OxmlElement("w:tblGrid")
        for cm in cm_widths:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(int(cm * CM)))
            tblGrid.append(gc)
        tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)

        for row_el in tbl.findall(qn("w:tr")):
            col_idx = 0
            for tc in row_el.findall(qn("w:tc")):
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    tc.insert(0, tcPr)
                gs  = tcPr.find(qn("w:gridSpan"))
                span = int(gs.get(qn("w:val"), 1)) if gs is not None else 1
                cell_twips = int(sum(
                    cm_widths[col_idx + k] * CM
                    for k in range(span) if col_idx + k < len(cm_widths)
                ))
                for old_w in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(old_w)
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"),    str(cell_twips))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)
                col_idx += span

    # ── Page break + title ────────────────────────────────────────────────────
    doc.add_page_break()

    title_p = doc.add_paragraph()
    title_r = title_p.add_run("Lampiran Penghitungan Imbalan Smart Payer")
    title_r.font.name  = body_font
    title_r.font.size  = Pt(body_size)
    title_r.font.bold  = True
    title_p.alignment  = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    rate = xlsx_data.get("daily_rate") or 0.00022
    try:    rate_pct = float(rate) * 100
    except: rate_pct = 0.022

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info_r = info_p.add_run(f"Imbalan Smart Payer = {rate_pct:.3f}% per hari")
    info_r.font.name = body_font
    info_r.font.size = Pt(body_size)
    info_r.font.bold = True

    doc.add_paragraph()

    # === Build the data table (unchanged from original) =======================
    rows_data       = xlsx_data.get("rows", [])
    date_columns    = xlsx_data.get("date_columns", [])
    totals          = xlsx_data.get("totals", {})
    pph_row         = xlsx_data.get("pph_row", {})
    net_disc_row    = xlsx_data.get("net_disc_row", {})
    transfer_row    = xlsx_data.get("transfer_row", {})
    n_date_cols     = len(date_columns)

    base_cols = 6
    total_cols = base_cols + n_date_cols
    n_rows = 2 + len(rows_data) + 1 + 3
    table = doc.add_table(rows=n_rows, cols=total_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    header_color = "1F4E78"
    subheader_color = "2E75B6"
    total_color = "BDD7EE"
    summary_color = "FFE699"

    # Row 0 — main headers
    hdr1 = table.rows[0]
    base_headers = ["Payer Name", "Document No", "Doc. Date", "Due Date", "Ostd.", "DPP"]
    for ci, txt in enumerate(base_headers):
        set_cell_text(hdr1.cells[ci], txt, bold=True, size=8,
                      color=(255, 255, 255),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr1.cells[ci], header_color)
        set_cell_margins(hdr1.cells[ci])

    if n_date_cols > 0:
        # Merge cells base_cols..total_cols-1 of hdr1 FIRST, then write the
        # "Tanggal Transfer..." label using write_merged_label so leftover
        # empty paragraphs from the merged-in cells get stripped.
        first_date_cell = hdr1.cells[base_cols]
        for ci in range(base_cols + 1, total_cols):
            first_date_cell.merge(hdr1.cells[ci])
        write_merged_label(first_date_cell,
                           "Tanggal Transfer / Tanggal Bayar - Reward Diberikan",
                           bold=True, size=8,
                           align=WD_ALIGN_PARAGRAPH.CENTER, shade=header_color)
        # Header text colour is white — write_merged_label uses set_cell_text
        # which doesn't set colour, so re-apply it directly to the run.
        for p in first_date_cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Row 1 — subheaders with dates
    hdr2 = table.rows[1]
    # Vertical merge of base-column headers (each spans rows 0..1). After
    # the merge, the cell contains the labelled paragraph from hdr1 AND an
    # empty paragraph carried over from hdr2 — strip the empty one so the
    # row doesn't double in height.
    for ci in range(base_cols):
        merged = hdr1.cells[ci].merge(hdr2.cells[ci])
        _tc = merged._tc
        # Drop empty trailing paragraphs (anything that has no <w:t> text).
        for _p in list(_tc.findall(qn("w:p"))):
            if not _p.findall(".//" + qn("w:t")):
                _tc.remove(_p)
        # Defensive: ensure at least one paragraph survives.
        if not _tc.findall(qn("w:p")):
            _tc.append(OxmlElement("w:p"))

    for di, dc in enumerate(date_columns):
        set_cell_text(hdr2.cells[base_cols + di], dc,
                      bold=True, size=7, color=(255, 255, 255),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr2.cells[base_cols + di], subheader_color)
        set_cell_margins(hdr2.cells[base_cols + di])

    # Data rows
    for ri, row_d in enumerate(rows_data):
        rr = table.rows[2 + ri]
        cells = rr.cells
        set_cell_text(cells[0], row_d["payer_name"], size=7,
                      align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(cells[1], row_d["doc_no"], size=7,
                      align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(cells[2], row_d["doc_date"], size=7,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[3], row_d["due_date"], size=7,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[4], fmt_num(row_d["ostd"]), size=7,
                      align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(cells[5], fmt_num(row_d["dpp"]), size=7,
                      align=WD_ALIGN_PARAGRAPH.RIGHT)
        for di, val in enumerate(row_d["rewards"]):
            set_cell_text(cells[base_cols + di], fmt_num(val), size=7,
                          align=WD_ALIGN_PARAGRAPH.RIGHT)
        for c in cells:
            set_cell_margins(c)

    # Totals row
    #
    # IMPORTANT: merge the four label cells FIRST, then write the "Total"
    # label into the resulting merged cell using write_merged_label() which
    # strips any leftover empty paragraphs. Writing empty text into cells
    # 1..3 and *then* merging (the original approach) leaves four extra
    # empty paragraphs inside the merged cell, inflating its rendered
    # height by ~4-5x.
    tot_row_idx = 2 + len(rows_data)
    tot = table.rows[tot_row_idx]
    # Merge cells 0..3 (Payer Name, Document No, Doc. Date, Due Date)
    tot.cells[0].merge(tot.cells[3])
    write_merged_label(tot.cells[0], "Total", bold=True, size=8,
                       align=WD_ALIGN_PARAGRAPH.CENTER, shade=total_color)
    set_cell_text(tot.cells[4], fmt_num(totals.get("ostd")), bold=True,
                  size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_text(tot.cells[5], fmt_num(totals.get("dpp")), bold=True,
                  size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
    reward_totals = totals.get("rewards", [])
    for di, val in enumerate(reward_totals):
        set_cell_text(tot.cells[base_cols + di], fmt_num(val),
                      bold=True, size=8,
                      align=WD_ALIGN_PARAGRAPH.RIGHT)
    # Shade + margins for the non-merged cells (cells 4, 5, and the date
    # columns). The merged cell was already shaded by write_merged_label.
    for c in list(tot.cells)[1:]:
        # tot.cells skips through gridSpan, but iterating the row's cells
        # still yields one entry per logical column post-merge — so the
        # first element after the merged label is cells[4]. Apply shade.
        shade_cell(c, total_color)
        set_cell_margins(c)

    # Summary rows — same merge-first pattern.
    summaries = [
        ("Total PPH 23",          pph_row.get("rewards", [])),
        ("Net Discount Diberikan", net_disc_row.get("rewards", [])),
        ("Nominal Transfer",       transfer_row.get("rewards", [])),
    ]
    for si, (label, vals) in enumerate(summaries):
        rr = table.rows[tot_row_idx + 1 + si]
        # Merge cells 0..5 (Payer Name through DPP) before writing the label.
        rr.cells[0].merge(rr.cells[base_cols - 1])
        write_merged_label(rr.cells[0], label, bold=True, size=8,
                           align=WD_ALIGN_PARAGRAPH.RIGHT, shade=summary_color)
        for di in range(n_date_cols):
            val = vals[di] if di < len(vals) else None
            set_cell_text(rr.cells[base_cols + di], fmt_num(val),
                          bold=True, size=8,
                          align=WD_ALIGN_PARAGRAPH.RIGHT)
        # Shade the un-merged date cells. The merged label cell was
        # already shaded by write_merged_label.
        for c in list(rr.cells)[1:]:
            shade_cell(c, summary_color)
            set_cell_margins(c)

    set_table_borders(table)

    # Column widths in cm. Tuned so the typical content fits on one line:
    #   Payer Name   1.6 cm    "BANLY THEO" / "PT. BAHAGIA INTRA NIAGA" wraps less
    #   Document No  1.4 cm    7-digit numbers fit comfortably
    #   Doc. Date    1.5 cm    "28/03/2026" (10 chars at 7pt) fits with margin
    #   Due Date     1.5 cm    same
    #   Ostd.        1.7 cm    "213,523,106" 11 chars
    #   DPP          2.0 cm    "192,363,158.56" 14 chars
    # Total base = 9.7 cm, leaving ~9.3 cm for the date columns (typically 4).
    base_widths_cm = [1.6, 1.4, 1.5, 1.5, 1.7, 2.0]
    base_total_cm  = sum(base_widths_cm)
    target_total_cm = 19.0
    if n_date_cols > 0:
        remaining_cm = max(target_total_cm - base_total_cm, n_date_cols * 1.5)
        date_cm = remaining_cm / n_date_cols
    else:
        date_cm = 1.0
    col_widths_cm = base_widths_cm + [date_cm] * n_date_cols
    apply_col_widths_cm(table, col_widths_cm)

    _set_table_page_margins(doc)

    # ── Convert team3/team4 email runs into mailto hyperlinks (NEW STEP) ─────
    if hyperlink_emails:
        _convert_emails_to_hyperlinks(doc, hyperlink_emails)

    doc.save(output_path)
    _print(f"  [OK] Tables appended: {output_path.name}")


# =============================================================================
# PDF CONVERSION
# =============================================================================

def convert_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """
    Convert DOCX → PDF using whichever office app is available (LibreOffice
    preferred for headless reliability; Word via COM as a fallback on Windows).
    """
    import shutil

    libre = shutil.which("soffice") or shutil.which("libreoffice")
    if not libre:
        for candidate in [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        ]:
            if Path(candidate).exists():
                libre = candidate
                break

    if libre:
        try:
            outdir = pdf_path.parent
            outdir.mkdir(parents=True, exist_ok=True)
            r = subprocess.run(
                [libre, "--headless", "--convert-to", "pdf",
                 "--outdir", str(outdir), str(docx_path)],
                capture_output=True, text=True, timeout=120,
            )
            produced = outdir / (docx_path.stem + ".pdf")
            if produced.exists():
                if produced != pdf_path:
                    try:
                        if pdf_path.exists():
                            pdf_path.unlink()
                        produced.rename(pdf_path)
                    except Exception:
                        pdf_path = produced
                _print(f"  [OK] PDF created: {pdf_path.name}")
                return True
            _print(f"  [!] LibreOffice did not produce a PDF.\n  STDOUT: {r.stdout}\n  STDERR: {r.stderr}")
            return False
        except Exception as e:
            _print(f"  [!] LibreOffice conversion error: {e}")

    if sys.platform == "win32":
        try:
            import comtypes.client
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            d = word.Documents.Open(str(docx_path.resolve()))
            d.SaveAs(str(pdf_path.resolve()), FileFormat=17)
            d.Close()
            word.Quit()
            _print(f"  [OK] PDF created via Word COM: {pdf_path.name}")
            return True
        except Exception as e:
            _print(f"  [!] PDF conversion error: {e}")
            return False


# =============================================================================
# MAIN PIPELINE
# =============================================================================

def run_pipeline(xlsx_path: Path,
                 use_defaults: bool = None,
                 batch_token: str | None = None,
                 auto_today: bool | None = None):
    """
    Execute the full pipeline for one XLSX file.

    Parameters
    ----------
    batch_token : a per-batch UUID supplied by the GUI/watcher. Used to
        derive the `nnn` sequence in `letter_number`. If None, the script
        falls back to a one-off batch (every run starts at 001).
    auto_today : if True (default when run from the GUI with the "Always
        use today's date" box ticked), overwrite the date-derived fields
        with values computed from datetime.now(). If False, trust whatever
        is in defaults / manual input.
    """
    _print("\n" + "=" * 65)
    _print("  Smart Payer Letter Generator")
    _print("  XLSX file:", xlsx_path.name)
    if batch_token:
        _print(f"  Batch token: {batch_token}")
    _print("=" * 65)

    defaults = load_defaults()

    # If the caller didn't tell us, fall back to whatever the defaults say.
    if auto_today is None:
        auto_today = bool(defaults.get("auto_today", False))

    # Step 1: determine variables
    if use_defaults is None:
        _print("\n  How would you like to fill the template variables?")
        _print("  [Y] Use default variables")
        _print("  [N] Enter variables manually")
        _print("  [S] Show current default variables")
        while True:
            try:
                choice = input("\n  Your choice (Y/N/S): ").strip().upper()
            except (EOFError, KeyboardInterrupt):
                choice = "Y"
            if choice == "S":
                show_defaults()
                continue
            if choice in ("Y", "N"):
                use_defaults = (choice == "Y")
                break

    if use_defaults:
        if not defaults:
            _print("\n  [!] No defaults found — switching to manual entry.")
            use_defaults = False
        else:
            variables = defaults.copy()
            _print("  [OK] Using default variables.")

    if not use_defaults:
        variables = prompt_all_variables(prefill=defaults)
        _print("\n  Save these as new defaults? (Y/N)")
        try:
            if input("  > ").strip().upper() == "Y":
                save_defaults(variables)
        except (EOFError, KeyboardInterrupt):
            pass

    # Strip schema-only metadata that should never reach the template engine.
    for k in ("auto_today", "auto_rates", "_comment_schema"):
        variables.pop(k, None)

    # ── Apply "Always use today's date" overrides (if active) ────────────────
    # These ALWAYS overwrite whatever was in defaults/user input — the whole
    # point of the toggle is that the user has surrendered manual control of
    # these fields.
    #
    # NOTE: month_name / month_year / end_of_month_date were historically
    # derived from today, but they are now ALWAYS derived from the XLSX
    # filename's encoded billing month (e.g. "Smartpayer April 2026 X.xlsx"
    # → month_year="April 2026", end_of_month_date="31 May 2026"). The
    # auto-today values for these three are only used as a fallback when
    # the filename doesn't carry a recognisable month/year.
    if auto_today:
        today_fields = compute_today_fields()
        for k in ("letter_date", "month_name", "month_year",
                  "month_roman", "year_num", "end_of_month_date"):
            variables[k] = today_fields[k]
        # Build the letter number: mmdd + zero-padded batch counter.
        counter = consume_batch_counter(batch_token)
        variables["letter_number"] = f"{today_fields['letter_number_prefix']}{counter:03d}"
        _print(f"  [auto-today] letter_date={variables['letter_date']}, "
               f"year_num={variables['year_num']}, "
               f"letter_number={variables['letter_number']}")
    else:
        # Even when auto-today is OFF, we still honour an explicit batch
        # token if present, so the user can opt into the dynamic nnn by
        # putting `mmdd` in their letter_number default (the trailing
        # three chars get replaced). This is best-effort — if their format
        # doesn't look like 7 digits we leave it alone.
        if batch_token and re.fullmatch(r"\d{7}", str(variables.get("letter_number", ""))):
            counter = consume_batch_counter(batch_token)
            base = str(variables["letter_number"])[:4]
            variables["letter_number"] = f"{base}{counter:03d}"

    # ── Override month_name / month_year / end_of_month_date from the XLSX
    #    filename's billing month. ALWAYS runs (independent of auto_today)
    #    because these fields describe the XLSX's data, not "today".
    #    Falls back to the previously-set values (defaults or auto-today)
    #    when the filename doesn't match the expected pattern.
    filename_fields = compute_filename_fields(xlsx_path.name)
    if filename_fields:
        for k in ("month_name", "month_year", "end_of_month_date"):
            variables[k] = filename_fields[k]
        _print(f"  [auto-filename] month_year={variables['month_year']}, "
               f"month_name={variables['month_name']}, "
               f"end_of_month_date={variables['end_of_month_date']}")
    else:
        _print("  [!] Filename doesn't encode a recognisable month/year — "
               "using fallback values for month_name / month_year / "
               "end_of_month_date.")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    stem = xlsx_path.stem

    # ── Read XLSX EARLY so we can derive discount_from / discount_until ──────
    # before filling the DOCX template. The original pipeline read the
    # workbook AFTER fill_template; we now read it first, derive the two
    # date fields, then fill, then reuse the parsed data for the tables.
    _print("\n  [Step 1/3] Parsing XLSX & deriving discount range...")
    xlsx_data = extract_xlsx_data(xlsx_path)
    d_from, d_until = derive_discount_range(xlsx_data)
    if d_from and d_until:
        variables["discount_from"]  = d_from
        variables["discount_until"] = d_until
        _print(f"  [auto-xlsx] discount_from={d_from}, discount_until={d_until}")
    else:
        _print("  [!] No datetime columns found in XLSX header — falling back "
               "to defaults for discount_from / discount_until.")

    # ── Derive the four reward-rate fields from XLSX daily_rate ───────────────
    # The cell immediately above "Simulasi Imbalan" (top-right of the
    # XLSX header) holds the daily reward rate as a decimal — e.g. 0.00022.
    # When auto-rates is on (default), use it to derive:
    #
    #   daily_rate            = 0.00022 × 100  -> "0.022"     (%/day)
    #   annual_rate           = 0.022 × 365    -> "8"         (%/year)
    #   rate_multiplier_num   = 8 / deposit_rate -> "4"
    #   rate_multiplier_word  = "empat"
    #
    # The auto_rates flag is persisted in defaults so users without the
    # daily_rate cell in their XLSX can opt out and keep manual control.
    auto_rates = bool(defaults.get("auto_rates", True))
    daily_decimal = xlsx_data.get("daily_rate")
    if auto_rates and daily_decimal is not None:
        derived = compute_rate_fields(daily_decimal, variables.get("deposit_rate"))
        if derived:
            variables.update(derived)
            _print(f"  [auto-rate] daily_rate={derived['daily_rate']}%/day, "
                   f"annual_rate={derived['annual_rate']}%/year, "
                   f"multiplier={derived['rate_multiplier_num']} ({derived['rate_multiplier_word']})")
        else:
            _print("  [!] Daily rate in XLSX is not a usable float — keeping defaults.")
    elif auto_rates:
        _print("  [!] No daily_rate cell found in XLSX header — keeping defaults.")

    # ── Fill DOCX template ────────────────────────────────────────────────────
    _print("\n  [Step 2/3] Filling DOCX template...")
    filled_docx = OUTPUT_DIR / f"{stem}_filled.docx"
    fill_template(variables, TEMPLATE_FILE, filled_docx)

    # ── Append tables + hyperlink team3/team4 emails ──────────────────────────
    _print("\n  [Step 2b/3] Appending tables + converting emails to hyperlinks...")
    final_docx = OUTPUT_DIR / f"{stem}_with_tables.docx"
    hyperlink_emails = [
        e for e in (variables.get("team3_email"), variables.get("team4_email"))
        if e
    ]
    append_tables_to_docx(filled_docx, xlsx_data, final_docx,
                          hyperlink_emails=hyperlink_emails)

    # ── PDF ──────────────────────────────────────────────────────────────────
    _print("\n  [Step 3/3] Converting to PDF...")
    pdf_path = OUTPUT_DIR / f"{stem}_letter.pdf"
    success = convert_to_pdf(final_docx, pdf_path)

    # Clean up intermediate file
    try:
        filled_docx.unlink(missing_ok=True)
    except Exception:
        pass

    _print("\n" + "=" * 65)
    if success:
        _print(f"  Done! ->  {pdf_path}")
    else:
        _print(f"  [!] PDF step failed. DOCX ready at: {final_docx}")
    _print("=" * 65 + "\n")

    return pdf_path if success else final_docx


# =============================================================================
# WATCH MODE
# =============================================================================

def watch_folder(folder: Path, use_defaults: bool = None):
    try:
        from watchdog.observers import Observer
        from watchdog.events import FileSystemEventHandler
    except ImportError:
        _print("  Installing watchdog...")
        subprocess.run([sys.executable, "-m", "pip", "install", "watchdog", "-q"],
                       check=True)
        from watchdog.observers import Observer
        from watchdog.events import FileSystemEventHandler

    import time
    processed = set()

    # CLI watch mode uses a single batch token for the whole watcher lifetime,
    # which means counters will keep climbing 001, 002, 003 ... rather than
    # resetting per scan-cycle. The GUI uses a smarter per-scan-cycle token.
    cli_batch_token = uuid.uuid4().hex

    class XLSXHandler(FileSystemEventHandler):
        def on_created(self, event):
            if event.is_directory:
                return
            path = Path(event.src_path)
            if path.suffix.lower() == ".xlsx" and path not in processed:
                processed.add(path)
                _print(f"\n  [NEW] File detected: {path.name}")
                import time as _t
                _t.sleep(1)
                run_pipeline(path, use_defaults=use_defaults,
                             batch_token=cli_batch_token)

    observer = Observer()
    observer.schedule(XLSXHandler(), str(folder), recursive=False)
    observer.start()
    _print(f"\n  [WATCH] Watching: {folder}")
    _print("  Waiting for new XLSX files... (Ctrl+C to stop)\n")
    try:
        while True:
            import time as _t
            _t.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()


# =============================================================================
# CLI ENTRY POINT
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description="Smart Payer Letter Generator")
    parser.add_argument("--xlsx",           help="Path to XLSX file to process")
    parser.add_argument("--watch",          help="Watch a folder for new XLSX files")
    parser.add_argument("--setup-defaults", action="store_true")
    parser.add_argument("--show-defaults",  action="store_true")
    parser.add_argument("--use-defaults",   action="store_true")
    parser.add_argument("--no-defaults",    action="store_true")
    parser.add_argument("--template",       help="Override DOCX template path")
    parser.add_argument("--output-dir",     help="Override output directory")
    # NEW: per-batch token + explicit auto-today flag (the GUI passes these)
    parser.add_argument("--batch-token",    help="Opaque per-batch identifier (UUID) used to drive the nnn counter")
    parser.add_argument("--reset-batch",    action="store_true",
                        help="Reset the batch state to counter=0 for --batch-token and exit (used by the GUI to claim a token)")
    auto_grp = parser.add_mutually_exclusive_group()
    auto_grp.add_argument("--auto-today",    dest="auto_today", action="store_true",
                          help="Force-enable 'Always use today's date'")
    auto_grp.add_argument("--no-auto-today", dest="auto_today", action="store_false",
                          help="Force-disable 'Always use today's date'")
    parser.set_defaults(auto_today=None)
    args = parser.parse_args()

    global TEMPLATE_FILE, OUTPUT_DIR
    if args.template:
        TEMPLATE_FILE = Path(args.template)
    if args.output_dir:
        OUTPUT_DIR = Path(args.output_dir)

    # GUI calls --reset-batch --batch-token <uuid> once per batch to zero
    # the counter before launching the first per-file subprocess.
    if args.reset_batch:
        if not args.batch_token:
            _print("[!] --reset-batch requires --batch-token")
            sys.exit(2)
        reset_batch_counter(args.batch_token)
        _print(f"[OK] Batch counter reset for token {args.batch_token}")
        return

    use_defaults = True if args.use_defaults else (False if args.no_defaults else None)

    if args.setup_defaults:
        setup_defaults()
    elif args.show_defaults:
        show_defaults()
    elif args.xlsx:
        run_pipeline(Path(args.xlsx),
                     use_defaults=use_defaults,
                     batch_token=args.batch_token,
                     auto_today=args.auto_today)
    elif args.watch:
        watch_folder(Path(args.watch), use_defaults=use_defaults)
    else:
        _print("\n=== Smart Payer Letter Generator ===")
        _print("  No arguments provided. Running in interactive mode.\n")
        while True:
            try:
                xlsx_input = input("  Enter path to XLSX file (or 'q' to quit): ").strip()
            except (EOFError, KeyboardInterrupt):
                break
            if xlsx_input.lower() == "q":
                break
            p = Path(xlsx_input)
            if not p.exists():
                _print(f"  File not found: {p}")
                continue
            run_pipeline(p, use_defaults=use_defaults,
                         batch_token=args.batch_token,
                         auto_today=args.auto_today)
            break


if __name__ == "__main__":
    main()
